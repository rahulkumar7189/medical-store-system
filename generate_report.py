import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Medical Store System - Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Author: Rahul Kumar').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Repository: https://github.com/rahulkumar7189/medical-store-system').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Project Overview
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "The Pharmacy Management System is a full-fledged desktop application built using JavaFX and MySQL. "
        "It provides a robust, user-friendly interface for managing pharmacy operations, including user authentication, "
        "inventory management, point of sale (billing), and supplier tracking. The application follows a strict "
        "Model-View-Controller (MVC) architecture, ensuring separation of concerns and maintainability."
    )

    # 2. Features
    doc.add_heading('2. Key Features', level=1)
    features = doc.add_paragraph()
    features.add_run("• User Authentication:").bold = True
    features.add_run(" Secure login system with role-based access.\n")
    features.add_run("• Master Dashboard:").bold = True
    features.add_run(" A sleek BorderPane layout enabling seamless view routing.\n")
    features.add_run("• Inventory Management:").bold = True
    features.add_run(" Real-time tracking of medicinal stock levels, categories, and expiration dates.\n")
    features.add_run("• Point of Sale (Billing):").bold = True
    features.add_run(" Interactive shopping cart that calculates subtotals, applies tax, processes sales, and deducts stock using atomic SQL Transactions.\n")
    features.add_run("• Supplier Directory:").bold = True
    features.add_run(" A module to manage pharmaceutical vendors and their contact information.")

    # 3. Technical Architecture
    doc.add_heading('3. Technical Architecture', level=1)
    arch = doc.add_paragraph()
    arch.add_run("• Language:").bold = True
    arch.add_run(" Java 17+\n")
    arch.add_run("• GUI Framework:").bold = True
    arch.add_run(" JavaFX (FXML + CSS)\n")
    arch.add_run("• Database:").bold = True
    arch.add_run(" MySQL Server\n")
    arch.add_run("• Build Tool:").bold = True
    arch.add_run(" Maven\n")
    arch.add_run("• Architecture Pattern:").bold = True
    arch.add_run(" MVC (Model-View-Controller)")

    # 4. Prompt Design & Interaction
    doc.add_heading('4. Prompt Design & Interaction Approach', level=1)
    doc.add_paragraph(
        "Throughout the development of this project, prompt engineering played a crucial role. "
        "The process started with defining core modules (Authentication, Inventory, POS, Admin) and designing "
        "a highly normalized database schema. Subsequent prompts guided the implementation of the Java POJOs, "
        "the Data Access Object (DAO) layer, and the JavaFX controllers. Clear boundaries were used to ensure "
        "code modularity and SQL transaction safety during the Point of Sale logic."
    )

    # 5. Challenges and Solutions
    doc.add_heading('5. Challenges and Solutions', level=1)
    
    doc.add_heading('Environment Setup & Port Discovery', level=2)
    doc.add_paragraph(
        "Challenge: The local MySQL server was configured to listen on a non-standard port (4000) instead of the default 3306, "
        "causing immediate connection refusal for the Java application.\n"
        "Solution: Performed system-wide port scanning and configuration file (my.ini) analysis to identify port 4000, "
        "then updated the application's .env configuration to align with the server environment."
    )

    doc.add_heading('Portable Build Environment', level=2)
    doc.add_paragraph(
        "Challenge: The target system lacked a globally configured Maven installation, preventing the standard build process.\n"
        "Solution: Implemented a portable Maven setup within the project workspace and updated the runner script (run.bat) "
        "to utilize this local binary, ensuring the project remains self-contained and runnable on machines without pre-installed build tools."
    )

    doc.add_heading('Database Seeding & Authentication', level=2)
    doc.add_paragraph(
        "Challenge: Initial application launch showed an empty Users table, leading to a hang on the authentication screen.\n"
        "Solution: Manually executed SQL seeding scripts to create the default 'admin' account, satisfying the application's "
        "authentication requirements and allowing access to the main dashboard."
    )

    doc.add_paragraph(
        "Challenge: Ensuring stock levels decrement accurately during concurrent or multi-item sales.\n"
        "Solution: Utilized strict SQL Transactions (connection.setAutoCommit(false)) in the SaleDao to ensure atomic inserts and rollbacks across the Sale, SaleItems, and Medicines tables.\n\n"
        "Challenge: Switching application views cleanly without creating multiple separate OS windows.\n"
        "Solution: Implemented a central BorderPane structure in the Main Dashboard that dynamically swaps its center StackPane node using FXMLLoader."
    )

    doc.add_page_break()

    # 6. Source Code Snippets
    doc.add_heading('6. Core Source Code Snippets', level=1)

    # Snippet 1: Database Connection
    doc.add_heading('DatabaseConnection.java (Snippet)', level=2)
    code_db = """public class DatabaseConnection {
    private static Connection connection = null;

    private DatabaseConnection() {}

    public static Connection getConnection() {
        if (connection == null) {
            try {
                Dotenv dotenv = Dotenv.load();
                String url = "jdbc:mysql://" + dotenv.get("DB_HOST") + ":" + dotenv.get("DB_PORT") + "/pharmacy_db";
                connection = DriverManager.getConnection(url, dotenv.get("DB_USER"), dotenv.get("DB_PASSWORD"));
            } catch (SQLException e) {
                e.printStackTrace();
            }
        }
        return connection;
    }
}"""
    p1 = doc.add_paragraph(code_db)
    p1.style = 'Normal'
    for run in p1.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    # Snippet 2: Dashboard View Routing
    doc.add_heading('DashboardController.java (Snippet)', level=2)
    code_route = """private void loadView(String fxmlPath) {
    try {
        FXMLLoader loader = new FXMLLoader(getClass().getResource(fxmlPath));
        Parent view = loader.load();
        Object controller = loader.getController();
        
        // Pass user context if it's the POS Controller
        if (controller instanceof POSController) {
            ((POSController) controller).setCurrentUser(loggedInUser);
        }

        contentArea.getChildren().clear();
        contentArea.getChildren().add(view);
    } catch (IOException e) {
        e.printStackTrace();
    }
}"""
    p2 = doc.add_paragraph(code_route)
    p2.style = 'Normal'
    for run in p2.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        
    # Snippet 3: POS Checking Out (SaleDao)
    doc.add_heading('SaleDao.java (Snippet)', level=2)
    code_pos = """public boolean processSale(Sale sale, List<SaleItem> items) {
    Connection conn = DatabaseConnection.getConnection();
    try {
        conn.setAutoCommit(false);
        // Insert Sale Record
        // Insert Loop for SaleItems
        // Update Stock Quantities
        conn.commit();
        return true;
    } catch (SQLException e) {
        try { conn.rollback(); } catch (SQLException ex) { }
        return false;
    } finally {
        try { conn.setAutoCommit(true); } catch (SQLException e) { }
    }
}"""
    p3 = doc.add_paragraph(code_pos)
    p3.style = 'Normal'
    for run in p3.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    output_path = "Medical_Store_System_Project_Report.docx"
    doc.save(output_path)
    print(f"Report fully generated successfully at: {output_path}")

if __name__ == "__main__":
    main()
